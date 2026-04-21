"""
EDRI — Document Processor (Metoda 3)
Parsare nativă + AI fallback
Fișier NOU — nu modifică nimic existent
"""

import os
import json
import re
from pathlib import Path

# ══ DETECTARE FORMAT ══
def detect_format(filename, file_bytes=None):
    ext = Path(filename).suffix.lower()
    fmt_map = {
        '.xlsx': 'xlsx', '.xls': 'xls',
        '.pdf': 'pdf',
        '.docx': 'docx', '.doc': 'doc',
        '.html': 'html', '.htm': 'html',
        '.csv': 'csv',
        '.jpg': 'image', '.jpeg': 'image',
        '.png': 'image', '.tiff': 'image'
    }
    return fmt_map.get(ext, 'unknown')

# ══ PROCESARE XLSX / XLS — parsare nativa ══
def process_excel(file_path):
    """Parsare nativa Excel cu openpyxl — precizie maxima, zero pierdere date"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        result = {}
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                clean_row = [str(cell) if cell is not None else '' for cell in row]
                if any(c.strip() for c in clean_row):
                    rows.append(clean_row)
            result[sheet_name] = rows
        return {'method': 'openpyxl_native', 'data': result, 'status': 'ok'}
    except ImportError:
        return process_excel_xlrd(file_path)
    except Exception as e:
        return {'method': 'openpyxl_native', 'data': {}, 'status': 'error', 'error': str(e)}

def process_excel_xlrd(file_path):
    """Fallback pentru .xls vechi"""
    try:
        import xlrd
        wb = xlrd.open_workbook(file_path)
        result = {}
        for sheet in wb.sheets():
            rows = []
            for i in range(sheet.nrows):
                rows.append([str(sheet.cell_value(i,j)) for j in range(sheet.ncols)])
            result[sheet.name] = rows
        return {'method': 'xlrd_native', 'data': result, 'status': 'ok'}
    except Exception as e:
        return {'method': 'xlrd_native', 'data': {}, 'status': 'error', 'error': str(e)}

# ══ PROCESARE PDF TEXT — parsare nativa ══
def process_pdf_text(file_path):
    """Parsare PDF text cu pdfplumber — pentru PDF-uri cu text selectabil"""
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ''
                tables = page.extract_tables() or []
                pages.append({
                    'page': i + 1,
                    'text': text,
                    'tables': tables
                })
        return {'method': 'pdfplumber_native', 'data': pages, 'status': 'ok'}
    except Exception as e:
        return {'method': 'pdfplumber_native', 'data': [], 'status': 'error', 'error': str(e)}

# ══ PROCESARE PDF SCANAT / IMAGINI — AI fallback ══
def process_with_ai(file_path, file_type='pdf_scanned'):
    """
    AI Document Intelligence — Azure fallback pentru:
    - PDF-uri scanate (fara text selectabil)
    - Imagini JPEG/PNG
    NOTA: Necesita AZURE_DOC_INTEL_KEY in environment variables
    """
    azure_key = os.environ.get('AZURE_DOC_INTEL_KEY')
    azure_endpoint = os.environ.get('AZURE_DOC_INTEL_ENDPOINT')

    if not azure_key or not azure_endpoint:
        # Fallback la pytesseract daca Azure nu e configurat
        return process_with_tesseract(file_path)

    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.core.credentials import AzureKeyCredential

        client = DocumentIntelligenceClient(
            endpoint=azure_endpoint,
            credential=AzureKeyCredential(azure_key)
        )
        with open(file_path, 'rb') as f:
            poller = client.begin_analyze_document(
                model_id='prebuilt-document',
                analyze_request=f,
                content_type='application/octet-stream'
            )
        result = poller.result()
        tables = []
        for table in result.tables or []:
            t = []
            for cell in table.cells:
                t.append({'row': cell.row_index, 'col': cell.column_index, 'content': cell.content})
            tables.append(t)
        return {
            'method': 'azure_document_intelligence',
            'data': {'content': result.content, 'tables': tables},
            'status': 'ok'
        }
    except Exception as e:
        return {'method': 'azure_ai', 'data': {}, 'status': 'error', 'error': str(e)}

def process_with_tesseract(file_path):
    """Fallback gratuit cu pytesseract daca Azure nu e disponibil"""
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(file_path)
        text = pytesseract.image_to_string(img, lang='ron+eng')
        return {'method': 'tesseract_ocr', 'data': {'text': text}, 'status': 'ok'}
    except Exception as e:
        return {'method': 'tesseract_ocr', 'data': {}, 'status': 'error', 'error': str(e)}

# ══ PROCESARE DOCX ══
def process_docx(file_path):
    """Parsare Word cu python-docx"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            t = []
            for row in table.rows:
                t.append([cell.text for cell in row.cells])
            tables.append(t)
        return {'method': 'python_docx_native', 'data': {'paragraphs': paragraphs, 'tables': tables}, 'status': 'ok'}
    except Exception as e:
        return {'method': 'python_docx_native', 'data': {}, 'status': 'error', 'error': str(e)}

# ══ PROCESARE HTML ══
def process_html(file_path):
    """Parsare HTML cu BeautifulSoup"""
    try:
        from bs4 import BeautifulSoup
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        tables = []
        for table in soup.find_all('table'):
            t = []
            for row in table.find_all('tr'):
                t.append([cell.get_text(strip=True) for cell in row.find_all(['td','th'])])
            tables.append(t)
        return {'method': 'beautifulsoup_native', 'data': {'tables': tables, 'text': soup.get_text()}, 'status': 'ok'}
    except Exception as e:
        return {'method': 'beautifulsoup_native', 'data': {}, 'status': 'error', 'error': str(e)}

# ══ PROCESARE CSV ══
def process_csv(file_path):
    """Parsare CSV cu pandas"""
    try:
        import pandas as pd
        df = pd.read_csv(file_path, encoding='utf-8', errors='replace')
        return {'method': 'pandas_native', 'data': df.to_dict(orient='records'), 'status': 'ok'}
    except Exception as e:
        return {'method': 'pandas_native', 'data': [], 'status': 'error', 'error': str(e)}

# ══ DETECTARE DACA PDF ARE TEXT SAU E SCANAT ══
def is_pdf_scanned(file_path):
    """Verifica daca un PDF contine text selectabil sau e scanat"""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages[:3]:
                text = page.extract_text() or ''
                if len(text.strip()) > 50:
                    return False
        return True
    except:
        return True

# ══ METODA 3 — ROUTER PRINCIPAL ══
def process_document(file_path, filename=None):
    """
    METODA 3 — Parsare nativă + AI fallback
    Router principal care alege metoda corectă automat
    """
    filename = filename or os.path.basename(file_path)
    fmt = detect_format(filename)

    if fmt in ('xlsx', 'xls'):
        result = process_excel(file_path)

    elif fmt == 'pdf':
        if is_pdf_scanned(file_path):
            result = process_with_ai(file_path, 'pdf_scanned')
        else:
            result = process_pdf_text(file_path)

    elif fmt in ('docx', 'doc'):
        result = process_docx(file_path)

    elif fmt == 'html':
        result = process_html(file_path)

    elif fmt == 'csv':
        result = process_csv(file_path)

    elif fmt == 'image':
        result = process_with_ai(file_path, 'image')

    else:
        result = {'method': 'unknown', 'data': {}, 'status': 'error', 'error': f'Format nesuportat: {fmt}'}

    result['filename'] = filename
    result['format'] = fmt
    return result

# ══ MAPARE CONTURI CONTABILE → INDICATORI ══
CONTURI_MAP = {
    '707': 'venituri_vanzari',
    '607': 'cost_marfuri',
    '5121': 'disponibil_banca',
    '4111': 'creante_clienti',
    '401': 'datorii_furnizori',
    '121': 'profit_pierdere',
    '681': 'amortizare',
    '3xx': 'stocuri',
    '411': 'creante',
}

def extract_financial_indicators(raw_data):
    """
    Extrage indicatorii financiari din datele brute parsate
    Cauta solduri finale pentru conturile contabile relevante
    """
    indicators = {}
    text_all = json.dumps(raw_data).upper()

    # Cauta solduri pe conturi
    for cont, indicator in CONTURI_MAP.items():
        pattern = rf'{cont}[^\d]*(\d[\d\.,]+)'
        matches = re.findall(pattern, text_all)
        if matches:
            try:
                val = float(matches[-1].replace('.', '').replace(',', '.'))
                indicators[indicator] = val
            except:
                pass

    # Calcul indicatori derivati
    v = indicators.get('venituri_vanzari', 0)
    c = indicators.get('cost_marfuri', 0)
    cash = indicators.get('disponibil_banca', 0)
    creante = indicators.get('creante_clienti', 0)
    profit = indicators.get('profit_pierdere', 0)

    if v > 0 and c > 0:
        indicators['marja_bruta_pct'] = round((v - c) / v * 100, 2)

    if cash > 0:
        cheltuieli_zilnice = (c / 90) if c > 0 else 1
        indicators['cash_runway_zile'] = round(cash / cheltuieli_zilnice)

    if creante > 0 and v > 0:
        indicators['dso_zile'] = round(creante / (v / 365))

    if profit > 0:
        indicators['ebitda'] = profit

    return indicators

if __name__ == '__main__':
    print("EDRI Document Processor — Metoda 3 — OK")
    print("Formate suportate: XLSX, XLS, PDF, DOCX, HTML, CSV, JPEG, PNG")
